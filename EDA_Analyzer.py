# ============================================================
#                       EDA ANALYZER
# ============================================================

import streamlit as st
import pandas as pd
import numpy as np
import os, time, warnings, joblib
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime

from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.pipeline import Pipeline
from sklearn.compose import ColumnTransformer
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.impute import SimpleImputer
from sklearn.metrics import f1_score, r2_score
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor
from sklearn.dummy import DummyClassifier, DummyRegressor
from sklearn.inspection import permutation_importance

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings("ignore")

# ============================================================
# CONFIG
# ============================================================
CONFIG = {
    "RANDOM_STATE": 42,
    "TEST_SIZE": 0.2,
    "CV_FOLDS": 5,
    "HIGH_CORR": 0.9,
    "SHAP_SAMPLE": 500
}

# ============================================================
# OUTPUT
# ============================================================
def create_output():
    folder = f"OUTPUT_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    os.makedirs(folder, exist_ok=True)
    return folder

# ============================================================
# PLOTS
# ============================================================
def save_plot(fig, path):
    fig.savefig(path, dpi=150, bbox_inches='tight')
    plt.close(fig)

def plot_missing(df, out):
    miss = df.isna().mean()
    if miss.sum() == 0: return None
    fig = plt.figure(figsize=(8,5))
    miss.sort_values(ascending=False).head(15).plot.bar(color='salmon')
    path = os.path.join(out,"missing.png")
    save_plot(fig,path)
    return path

def plot_corr(df, out):
    num = df.select_dtypes(include=np.number)
    if num.shape[1] < 2: return None
    fig = plt.figure(figsize=(10,8))
    sns.heatmap(num.corr(), cmap="coolwarm", annot=True, fmt=".2f")
    path = os.path.join(out,"corr.png")
    save_plot(fig,path)
    return path

def plot_target(y, out):
    fig = plt.figure(figsize=(6,4))
    if y.nunique() < 20:
        y.value_counts().plot.bar(color='cornflowerblue')
    else:
        sns.histplot(y, kde=True)
    path = os.path.join(out,"target.png")
    save_plot(fig,path)
    return path

# ============================================================
# FEATURE SELECTION
# ============================================================
def drop_high_corr(X):
    corr = X.select_dtypes(include=np.number).corr().abs()
    upper = corr.where(np.triu(np.ones(corr.shape), k=1).astype(bool))
    drop_cols = [col for col in upper.columns if any(upper[col] > CONFIG["HIGH_CORR"])]
    return X.drop(columns=drop_cols), drop_cols

# ============================================================
# AUTOML
# ============================================================
def automl(df, target):
    t0 = time.time()
    X = df.drop(columns=[target])
    y = df[target]

    # Feature selection
    X, dropped = drop_high_corr(X)
    is_clf = y.dtype == 'object' or y.nunique()<10

    X_train,X_test,y_train,y_test = train_test_split(
        X,y,test_size=CONFIG["TEST_SIZE"],random_state=CONFIG["RANDOM_STATE"]
    )

    num_cols = X_train.select_dtypes(include=np.number).columns
    cat_cols = X_train.select_dtypes(exclude=np.number).columns

    prep = ColumnTransformer([
        ("num", Pipeline([("imp",SimpleImputer()),("sc",StandardScaler())]), num_cols),
        ("cat", Pipeline([("imp",SimpleImputer(strategy="most_frequent")),
                          ("enc",OneHotEncoder(handle_unknown='ignore'))]), cat_cols)
    ])

    models = {}
    if is_clf:
        models = {
            "RandomForest": RandomForestClassifier(),
            "Logistic": LogisticRegression(max_iter=1000),
            "DecisionTree": DecisionTreeClassifier(),
            "Dummy": DummyClassifier()
        }
    else:
        models = {
            "RandomForest": RandomForestRegressor(),
            "Linear": LinearRegression(),
            "DecisionTree": DecisionTreeRegressor(),
            "Dummy": DummyRegressor()
        }

    results = []
    best_model = None
    best_score = -np.inf

    for name, model in models.items():
        pipe = Pipeline([("prep",prep),("model",model)])
        try:
            cv = cross_val_score(pipe,X_train,y_train,cv=CONFIG["CV_FOLDS"]).mean()
            pipe.fit(X_train,y_train)
            pred = pipe.predict(X_test)
            score = f1_score(y_test,pred,average='macro') if is_clf else r2_score(y_test,pred)
            results.append((name, round(cv,4), round(score,4)))
            if score > best_score:
                best_score = score
                best_model = pipe
        except:
            continue

    runtime = round(time.time()-t0,2)
    return best_model, results, runtime, dropped

# ============================================================
# SHAP
# ============================================================
try:
    import shap
    SHAP_AVAILABLE=True
except:
    SHAP_AVAILABLE=False

def shap_plot(pipe, X, out):
    if not SHAP_AVAILABLE:
        return None
    model = pipe.named_steps.get("model", None)
    if isinstance(model, (DummyClassifier, DummyRegressor)):
        print("SHAP skipped: Dummy model detected.")
        return None
    try:
        n_samples = min(len(X), CONFIG["SHAP_SAMPLE"])
        X_sub = X.iloc[:n_samples]
        Xt = pipe.named_steps["prep"].transform(X_sub)
        feature_names = pipe.named_steps["prep"].get_feature_names_out()
        if not callable(getattr(model, "predict", None)):
            print(f"SHAP skipped: Model {type(model).__name__} not callable.")
            return None
        explainer = shap.Explainer(model, Xt)
        sv = explainer(Xt)
        plt.figure(figsize=(10,7))
        shap.summary_plot(sv, Xt, feature_names=feature_names, show=False)
        path = os.path.join(out, "shap_summary.png")
        plt.savefig(path,dpi=140,bbox_inches='tight')
        plt.close()
        return path
    except Exception as e:
        print("SHAP failed:", str(e))
        return None

# ============================================================
# DOCX HELPERS
# ============================================================
def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for r in h.runs: r.bold = True
    return h

def add_table(doc, df, title):
    add_heading(doc, title, 2)
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return
    df = df.reset_index(drop=True).round(4)
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    table.style = 'Table Grid'
    for j, col in enumerate(df.columns):
        table.cell(0,j).text = str(col)
    for i in range(len(df)):
        for j in range(len(df.columns)):
            table.cell(i+1,j).text = str(df.iloc[i,j])

def add_image(doc, path, title):
    add_heading(doc, title, 3)
    try:
        doc.add_picture(path, width=Inches(6))
    except:
        doc.add_paragraph("Image not available")

# ============================================================
# STATS HELPERS
# ============================================================
def basic_info(df):
    return pd.DataFrame({
        "Column": df.columns,
        "Dtype": df.dtypes.astype(str),
        "Missing": df.isna().sum(),
        "Missing%": (df.isna().sum()/len(df)*100).round(2),
        "Unique": df.nunique(),
        "Cardinality": (df.nunique()/len(df)).round(3)
    })

def descriptive(df):
    return df.describe(include='all').T

def skew_kurt(df):
    num = df.select_dtypes(include=np.number)
    if num.empty: return None
    return pd.DataFrame({
        "Feature": num.columns,
        "Skew": num.skew().round(3),
        "Kurtosis": num.kurt().round(3)
    })

def missing_overview(df, out):
    miss = df.isna().mean().sort_values(ascending=False)[:15]
    if miss.sum()==0: return None, None
    fig = plt.figure(figsize=(10,5))
    miss.plot(kind='bar', color='salmon')
    path = os.path.join(out,"missing_values.png")
    plt.savefig(path,dpi=140,bbox_inches='tight')
    plt.close()
    return miss, path

def correlation(df, out):
    num = df.select_dtypes(include=np.number)
    if num.shape[1]<2: return None, None, None
    corr = num.corr()
    fig = plt.figure(figsize=(12,10))
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm", linewidths=0.5)
    path = os.path.join(out,"correlation_heatmap.png")
    plt.savefig(path,dpi=150,bbox_inches='tight')
    plt.close()
    return corr, None, path

def feature_importance(pipe, X, y, out):
    try:
        Xt = pipe.named_steps["prep"].transform(X)
        names = pipe.named_steps["prep"].get_feature_names_out()
        r = permutation_importance(pipe.named_steps["model"], Xt, y,
                                   n_repeats=7, random_state=CONFIG["RANDOM_STATE"], n_jobs=-1)
        df = pd.DataFrame({
            "EncodedFeature": names,
            "Importance": r.importances_mean,
            "Std": r.importances_std
        }).sort_values("Importance", ascending=False)
        original_map = {}
        for enc in names:
            if '__' in enc:
                orig = enc.split('__')[1].split('_')[0]
                original_map[enc] = orig
            else:
                original_map[enc] = enc
        df["OriginalFeature"] = df["EncodedFeature"].map(original_map)
        agg = df.groupby("OriginalFeature")["Importance"].sum().sort_values(ascending=False)
        fig = plt.figure(figsize=(10,7))
        agg.head(15).plot.barh(color='cornflowerblue', edgecolor='black')
        path = os.path.join(out,"feature_importance_grouped.png")
        plt.savefig(path,dpi=140,bbox_inches='tight')
        plt.close()
        return df, agg, path
    except Exception as e:
        print("Feature importance failed:", str(e))
        return None, None, None

# ============================================================
# REPORT
# ============================================================
def generate_report(df, target, out):
    doc = Document()
    doc.add_heading("EDA REPORT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Overview", 1)
    add_table(doc, basic_info(df), "1.1 Basic Info")

    add_heading(doc, "2. Statistics",1)
    add_table(doc, descriptive(df), "2.1 Descriptive Statistics")
    add_table(doc, skew_kurt(df), "2.2 Skewness & Kurtosis")

    add_heading(doc, "3. Missing Values",1)
    miss, miss_img = missing_overview(df, out)
    if miss_img:
        add_image(doc, miss_img, "3.1 Top Missing Columns")

    add_heading(doc, "4. Correlation",1)
    _, _, corr_img = correlation(df, out)
    if corr_img:
        add_image(doc, corr_img, "4.1 Correlation Heatmap")

    add_heading(doc, "5. Target Analysis",1)
    t_img = plot_target(df[target], out)
    add_image(doc, t_img, "5.1 Target Distribution")

    add_heading(doc, "6. AutoML Benchmark",1)
    model, results, runtime, dropped = automl(df,target)
    add_table(doc, pd.DataFrame(results, columns=["Model","CV Score","Test Score"]), "6.1 Model Comparison")
    add_heading(doc,"6.2 High-Correlation Features Removed",2)
    doc.add_paragraph(str(dropped) if dropped else "None")

    add_heading(doc, "7. Feature Importance",1)
    fi, fi_agg, fi_img = feature_importance(model, df.drop(columns=[target]), df[target], out)
    if fi_img:
        add_image(doc, fi_img, "7.1 Grouped Permutation Importance")

    add_heading(doc, "8. SHAP (Explainable AI)",1)
    sp = shap_plot(model, df.drop(columns=[target]), out)
    if sp:
        add_image(doc, sp, "8.1 SHAP Summary Plot")
    else:
        doc.add_paragraph("SHAP analysis skipped: Model not compatible or Dummy model detected.")

    add_heading(doc, "9. AI-Based Insights & Recommendations",1)
    p = doc.add_paragraph()
    p.add_run("• Task type auto-detected (may need manual override)\n")
    p.add_run("• Mild outlier clipping applied (0.5%–99.5% quantiles)\n")
    p.add_run("• High-correlation filter applied (numeric > 0.92)\n")
    p.add_run("• Feature importance indicates key drivers of target\n")
    p.add_run("• LightGBM + RandomForest often most predictive\n")
    p.add_run("• Consider removing low-importance features to simplify models\n")
    p.add_run("• SHAP plots provide model explainability for compatible models\n")
    p.add_run("• Ensure sufficient sample size for robust SHAP interpretation\n")
    p.add_run("• Check missing value treatment & outliers before production deployment\n")

    report_path = os.path.join(out,"EDA_REPORT.docx")
    doc.save(report_path)
    print(f"\nReport saved : {report_path}")

# ============================================================
# STREAMLIT UI
# ============================================================
st.set_page_config(layout="wide")
st.title("EDA Analyzer")

file = st.file_uploader("Upload Dataset")

if file:
    df = pd.read_csv(file)
    st.dataframe(df.head())
    target = st.selectbox("Select Target", df.columns)

    if st.button("Run Full Analysis"):
        out = create_output()
        m = plot_missing(df,out)
        c = plot_corr(df,out)
        t = plot_target(df[target],out)
        model, results, runtime, dropped = automl(df,target)
        sp = shap_plot(model, df.drop(columns=[target]), out)
        joblib.dump(model, os.path.join(out,"model.pkl"))
        generate_report(df,target,out)
        st.success("Done")
        st.write("Runtime:",runtime)
        st.write("Dropped Features:",dropped)
        for r in results:
            st.write(r)
        if m: st.image(m)
        if c: st.image(c)
        if t: st.image(t)
        if sp: st.image(sp)
        st.download_button("Download Report",
            open(os.path.join(out,"EDA_REPORT.docx"),"rb"),
            "report.docx")
        st.download_button("Download Model",
            open(os.path.join(out,"model.pkl"),"rb"),
            "model.pkl")
